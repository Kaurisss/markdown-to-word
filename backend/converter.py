"""Markdown-to-Word document converter."""

import os
import re
from typing import Any, Dict

from docx import Document

from .errors import (
    FileError, PermissionError_, ConfigError,
    ConversionError, DocxGenerationError,
)
from .converters.code_block import flush_code_buffer, is_fence, process_code_buffer
from .converters.table import flush_table_buffer, process_table_buffer
from .converters.image import add_image
from .converters.toc import add_toc
from .document_layout import apply_document_layout
from .elements import (
    add_heading, add_body, add_caption, add_quote, add_list_item,
    add_horizontal_rule,
    set_page_margins,
)


def _is_output_permission_error(error: Exception) -> bool:
    """Check if an exception is caused by output file permission/lock issues."""
    if isinstance(error, PermissionError):
        return True

    winerror = getattr(error, "winerror", None)
    if winerror in {5, 32}:
        return True

    errno = getattr(error, "errno", None)
    if errno in {13, 16}:
        return True

    text = str(error).lower()
    locked_markers = (
        "permission denied",
        "being used by another process",
        "另一个程序正在使用此文件",
        "拒绝访问",
    )
    return any(marker in text for marker in locked_markers)


def convert(input_path: str, output_path: str, conf: Dict[str, Any], resource_root: str | None = None) -> None:
    """Convert Markdown file to Word document with proper error handling."""
    if not os.path.exists(input_path):
        raise FileError(
            "Input file not found",
            path=input_path
        )

    resource_root = resource_root or os.path.dirname(os.path.abspath(input_path))

    output_dir = os.path.dirname(output_path) or '.'
    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir, exist_ok=True)
        except PermissionError as e:
            raise PermissionError_(
                "Cannot create output directory",
                path=output_dir,
                details=str(e)
            )

    try:
        doc = Document()
    except Exception as e:
        raise DocxGenerationError(
            "Failed to create Word document",
            details=str(e)
        )

    try:
        margin_value = conf.get("global", {}).get("pageMargin", 1.0)
        if not isinstance(margin_value, (float, int, dict)):
            margin_value = float(margin_value)
    except Exception as e:
        raise ConfigError("Invalid pageMargin value", details=str(e))
    set_page_margins(doc, margin_value)

    add_toc(doc, conf)

    try:
        with open(input_path, "r", encoding="utf-8") as f:
            lines = f.read().splitlines()
    except PermissionError as e:
        raise PermissionError_(
            "Permission denied reading input file",
            path=input_path,
            details=str(e)
        )
    except UnicodeDecodeError as e:
        raise ConversionError(
            "Failed to decode input file (expected UTF-8 encoding)",
            details=str(e)
        )

    in_code = False
    code_buf: list[str] = []
    in_table = False
    table_buf: list[str] = []
    seen_heading = False
    caption_counts = {"图": 0, "表": 0, "公式": 0}
    image_caption_count = 0

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Flush pending table before a fence so document order is preserved.
        if is_fence(line) and in_table:
            flush_table_buffer(doc, table_buf, conf)
            table_buf = []
            in_table = False

        # Code block detection and buffering
        consumed, code_buf, in_code = process_code_buffer(
            doc, line, in_code, code_buf, conf,
        )
        if consumed:
            i += 1
            continue

        # Table detection and buffering
        next_line = lines[i + 1] if i + 1 < len(lines) else None
        consumed, table_buf, in_table = process_table_buffer(
            doc, line, in_table, table_buf, conf, next_line,
        )
        if consumed:
            i += 1
            continue

        image_match = re.match(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
        if image_match:
            alt_text, image_reference = image_match.groups()
            add_image(doc, image_reference.strip(), alt_text, resource_root, conf)
            image_config = conf.get("imageCaption", {})
            if image_config.get("useAltText") and alt_text.strip():
                image_caption_count += 1
                caption = alt_text.strip()
                if image_config.get("autoNumber"):
                    caption = f"图{image_caption_count} {caption}"
                add_caption(doc, caption, conf)
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            body_start = conf.get("global", {}).get("bodyStart", {})
            is_document_title = (
                bool(body_start.get("firstHeadingAsTitle", False))
                and not seen_heading
                and level == 1
            )
            if is_document_title:
                add_heading(doc, text, level, conf, is_document_title=True)
            else:
                add_heading(doc, text, level, conf)
            seen_heading = True
            i += 1
            continue
        numbered_caption_match = re.match(r"^\s*(图|表|公式)\s*(\d+)[\s　].+", line)
        if numbered_caption_match:
            kind = numbered_caption_match.group(1)
            caption_counts[kind] = max(caption_counts[kind], int(numbered_caption_match.group(2)))
            add_caption(doc, line.strip(), conf)
            i += 1
            continue
        auto_caption_match = re.match(r"^\s*(图|表|公式)[\s　]+(.+)$", line)
        if auto_caption_match:
            kind = auto_caption_match.group(1)
            caption_counts[kind] += 1
            add_caption(doc, f"{kind}{caption_counts[kind]} {auto_caption_match.group(2).strip()}", conf)
            i += 1
            continue
        if re.match(r"^\s*>\s+(.*)$", line):
            text = re.sub(r"^\s*>\s+", "", line)
            add_quote(doc, text, conf)
            i += 1
            continue
        if re.match(r"^\s*([-*_])\s*\1\s*\1\s*$", line.strip()):
            add_horizontal_rule(doc, conf)
            i += 1
            continue
        ul_match = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if ul_match:
            indent = ul_match.group(1)
            text = ul_match.group(2)
            indent_len = len(indent.replace('\t', '  '))
            level = indent_len // 2
            add_list_item(doc, text, ordered=False, conf=conf, level=level)
            i += 1
            continue
        ol_match = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if ol_match:
            indent = ol_match.group(1)
            text = ol_match.group(2)
            indent_len = len(indent.replace('\t', '  '))
            level = indent_len // 2
            add_list_item(doc, text, ordered=True, conf=conf, level=level)
            i += 1
            continue
        if line.strip() == "":
            i += 1
            continue
        add_body(doc, line, conf)
        i += 1

    # Flush any remaining buffers
    flush_code_buffer(doc, code_buf, conf)
    flush_table_buffer(doc, table_buf, conf)

    apply_document_layout(doc, conf)

    try:
        doc.save(output_path)
    except Exception as e:
        if _is_output_permission_error(e):
            raise PermissionError_(
                "Cannot write output file",
                path=output_path,
                details="The target file may be open in Word/WPS or locked by another application."
            )

        raise DocxGenerationError(
            "Failed to save Word document",
            details=str(e)
        )
