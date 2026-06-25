"""Word page, section, header, and footer layout helpers."""

from typing import Any, Dict

from docx import Document
from docx.shared import Cm, Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .styling import _ensure_east_asia_font, _get_alignment, hex_to_rgb


def _length(value: Any):
    """Convert a numeric config value to a Word length.

    Numeric values are treated as inches for backward compatibility. Dict values
    can opt into centimetres with {"value": 2.8, "unit": "cm"}.
    """
    if isinstance(value, dict):
        raw = value.get("value", value.get("cm", value.get("in", 0)))
        unit = str(value.get("unit", "cm" if "cm" in value else "in")).lower()
        return Cm(float(raw)) if unit == "cm" else Inches(float(raw))
    return Inches(float(value))


def _page_dimension(value: Any, default_unit: str = "in"):
    if isinstance(value, dict):
        raw = value.get("value", value.get("cm", value.get("in", 0)))
        unit = str(value.get("unit", default_unit)).lower()
        return Cm(float(raw)) if unit == "cm" else Inches(float(raw))
    return Cm(float(value)) if default_unit == "cm" else Inches(float(value))


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _paragraph_alignment(paragraph, value: str | None) -> None:
    align = _get_alignment(value)
    if align is not None:
        paragraph.alignment = align


def _apply_run_style(run, style_config: Dict[str, Any], global_config: Dict[str, Any]) -> None:
    font_size = style_config.get("fontSize")
    if font_size is not None:
        run.font.size = Pt(float(font_size))
    if style_config.get("bold") is not None:
        run.bold = bool(style_config.get("bold"))
    color_hex = style_config.get("color")
    if color_hex:
        run.font.color.rgb = hex_to_rgb(str(color_hex))

    cn_font = style_config.get("fontFamily") or global_config.get("baseFontCn") or "SimSun"
    en_font = style_config.get("fontFamilyEn") or global_config.get("baseFontEn") or cn_font
    _ensure_east_asia_font(run, cn_font, en_font)


def _add_field_run(paragraph, instr: str, style_config: Dict[str, Any], global_config: Dict[str, Any]) -> None:
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instr} "

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr_text)
    run._r.append(end)
    _apply_run_style(run, style_config, global_config)


def _add_text_run(paragraph, text: str, style_config: Dict[str, Any], global_config: Dict[str, Any]) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    _apply_run_style(run, style_config, global_config)


def _add_page_number_text(paragraph, fmt: str, style_config: Dict[str, Any], global_config: Dict[str, Any]) -> None:
    i = 0
    while i < len(fmt):
        if fmt.startswith("{page}", i):
            _add_field_run(paragraph, "PAGE", style_config, global_config)
            i += len("{page}")
        elif fmt.startswith("{pages}", i):
            _add_field_run(paragraph, "NUMPAGES", style_config, global_config)
            i += len("{pages}")
        else:
            next_page = fmt.find("{page}", i)
            next_pages = fmt.find("{pages}", i)
            candidates = [pos for pos in (next_page, next_pages) if pos != -1]
            next_token = min(candidates) if candidates else len(fmt)
            _add_text_run(paragraph, fmt[i:next_token], style_config, global_config)
            i = next_token


def _set_page_number_start(section, start: int) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def _set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _apply_page_size(section, global_config: Dict[str, Any]) -> None:
    page_size = global_config.get("pageSize")
    if not isinstance(page_size, dict):
        return
    unit = str(page_size.get("unit", "in")).lower()
    if page_size.get("width") is not None:
        section.page_width = _page_dimension(page_size["width"], unit)
    if page_size.get("height") is not None:
        section.page_height = _page_dimension(page_size["height"], unit)


def _apply_margins(section, margin: Any) -> None:
    if isinstance(margin, dict):
        if "top" in margin:
            section.top_margin = _length(margin["top"])
        if "bottom" in margin:
            section.bottom_margin = _length(margin["bottom"])
        if "left" in margin:
            section.left_margin = _length(margin["left"])
        if "right" in margin:
            section.right_margin = _length(margin["right"])
        return

    value = _length(margin)
    section.top_margin = value
    section.bottom_margin = value
    section.left_margin = value
    section.right_margin = value


def _apply_header(section, header_config: Dict[str, Any], global_config: Dict[str, Any]) -> None:
    if not header_config.get("enabled", False):
        return
    if header_config.get("distance") is not None:
        section.header_distance = _length(header_config["distance"])

    paragraph = section.header.paragraphs[0]
    _clear_paragraph(paragraph)
    _paragraph_alignment(paragraph, header_config.get("alignment", "center"))
    _add_text_run(paragraph, str(header_config.get("text", "")), header_config, global_config)


def _apply_footer(section, footer_config: Dict[str, Any], global_config: Dict[str, Any], *, clear_only: bool = False) -> None:
    if footer_config.get("distance") is not None:
        section.footer_distance = _length(footer_config["distance"])

    paragraph = section.footer.paragraphs[0]
    _clear_paragraph(paragraph)

    if clear_only or not footer_config.get("enabled", False):
        return

    _paragraph_alignment(paragraph, footer_config.get("alignment", "center"))
    if footer_config.get("pageNumber", False):
        _add_page_number_text(
            paragraph,
            str(footer_config.get("format", "第{page}页（共{pages}页）")),
            footer_config,
            global_config,
        )
    else:
        _add_text_run(paragraph, str(footer_config.get("text", "")), footer_config, global_config)


def apply_document_layout(doc: Document, conf: Dict[str, Any]) -> None:
    """Apply document-level Word layout options from the style config."""
    if not hasattr(doc, "sections"):
        return

    global_config = conf.get("global", {})
    margin = global_config.get("pageMargin", 1.0)
    header_config = global_config.get("header") or {}
    footer_config = global_config.get("footer") or {}
    body_start_config = global_config.get("bodyStart") or {}

    for section in doc.sections:
        _apply_page_size(section, global_config)
        _apply_margins(section, margin)
        _apply_header(section, header_config, global_config)

    if footer_config:
        start_at_body = bool(footer_config.get("startAtBody", False))
        has_body_section = start_at_body and len(doc.sections) > 1
        for index, section in enumerate(doc.sections):
            if index > 0:
                section.footer.is_linked_to_previous = False
            clear_only = has_body_section and index == 0
            _apply_footer(section, footer_config, global_config, clear_only=clear_only)

    if body_start_config.get("restartPageNumberAfterToc") and len(doc.sections) > 1:
        start = int(body_start_config.get("pageNumberStart", 1) or 1)
        _set_page_number_start(doc.sections[1], start)

    _set_update_fields_on_open(doc)
