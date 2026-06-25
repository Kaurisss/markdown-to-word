"""Document element builders: heading, body, quote, list, code, hr, and re-exports."""

import re
from typing import Any, Dict

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .converters.styles import (
    apply_paragraph_fmt, apply_run_fmt, _ensure_east_asia_font,
    _set_run_shading, _get_alignment,
)
from .parser import parse_inline_formatting
from .text_normalization import normalize_fullwidth_punctuation


def add_hyperlink(paragraph, url: str, text: str, style_config: Dict[str, Any], global_config: Dict[str, Any]):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    run = paragraph.add_run(text)
    apply_run_fmt(run, style_config, global_config)

    run.font.color.rgb = RGBColor(0, 0, 255)
    run.underline = True

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)


def add_formatted_runs(paragraph, text: str, base_style: Dict[str, Any], global_config: Dict[str, Any], code_style: Dict[str, Any] = None) -> None:
    """Add text with inline formatting to a paragraph."""
    segments = parse_inline_formatting(text)

    for segment in segments:
        if segment['link']:
            add_hyperlink(paragraph, segment['link'], segment['text'], base_style, global_config)
        else:
            segment_text = segment['text']
            if global_config.get("normalizePunctuation") and not segment.get("code"):
                segment_text = normalize_fullwidth_punctuation(segment_text)

            run = paragraph.add_run(segment_text)

            if segment['code'] and code_style:
                apply_run_fmt(run, code_style, global_config)
            else:
                apply_run_fmt(run, base_style, global_config)

            if segment.get('bold'):
                run.bold = True
            if segment.get('italic'):
                run.italic = True
            if segment.get('underline'):
                run.underline = True
            if segment.get('strike'):
                run.font.strike = True
            if segment.get('code'):
                code_font = code_style.get('fontFamily', 'Courier New') if code_style else 'Courier New'
                run.font.name = code_font
                _ensure_east_asia_font(run, global_config.get('baseFontCn', 'SimSun'), code_font)
                if code_style:
                    bg_color = code_style.get('backgroundColor', '#F5F7F9')
                    _set_run_shading(run, bg_color)


def add_heading(doc: Document, text: str, level: int, conf: Dict[str, Any], is_document_title: bool = False) -> None:
    if is_document_title:
        style = conf["styles"].get("documentTitle", conf["styles"].get("h1", {}))
    else:
        style = conf["styles"].get(f"h{level}", conf["styles"]["h1"])

    heading_style_map = {
        1: "Heading 1",
        2: "Heading 2",
        3: "Heading 3",
        4: "Heading 4",
        5: "Heading 5",
        6: "Heading 6",
    }

    builtin_style = "Title" if is_document_title else heading_style_map.get(level, "Heading 1")
    try:
        p = doc.add_paragraph(style=builtin_style)
    except KeyError:
        p = doc.add_paragraph()

    apply_paragraph_fmt(p, style)
    code_style = conf["styles"].get("code", {})
    add_formatted_runs(p, text.strip(), style, conf["global"], code_style)


def add_caption(doc: Document, text: str, conf: Dict[str, Any]) -> None:
    style = conf["styles"].get("caption", conf["styles"].get("body", {})).copy()
    style["firstLineIndent"] = 0

    p = doc.add_paragraph()
    apply_paragraph_fmt(p, style)
    code_style = conf["styles"].get("code", {})
    add_formatted_runs(p, text.strip(), style, conf["global"], code_style)


def add_body(doc: Document, text: str, conf: Dict[str, Any]) -> None:
    style = conf["styles"]["body"]
    p = doc.add_paragraph()
    apply_paragraph_fmt(p, style)

    stripped = text.lstrip()
    if re.match(r"^(\*\*|\*|`|\[|!\[)", stripped):
        p.paragraph_format.first_line_indent = 0

    code_style = conf["styles"].get("code", {})
    add_formatted_runs(p, text, style, conf["global"], code_style)


def add_quote(doc: Document, text: str, conf: Dict[str, Any]) -> None:
    style = conf["styles"]["quote"]
    p = doc.add_paragraph()
    apply_paragraph_fmt(p, style)
    try:
        p.paragraph_format.left_indent = Inches(0.25)
    except Exception:
        pass
    code_style = conf["styles"].get("code", {})
    add_formatted_runs(p, text, style, conf["global"], code_style)


def add_list_item(doc: Document, text: str, ordered: bool, conf: Dict[str, Any], level: int = 0) -> None:
    style = conf["styles"]["body"].copy()
    style["firstLineIndent"] = 0

    p = doc.add_paragraph()
    apply_paragraph_fmt(p, style)

    try:
        p.style = "List Number" if ordered else "List Bullet"
    except Exception:
        pass

    if level > 0:
        try:
            p.paragraph_format.left_indent = Inches(level * 0.5)
        except Exception:
            pass

    code_style = conf["styles"].get("code", {})
    add_formatted_runs(p, text, style, conf["global"], code_style)


def add_code_block(doc: Document, lines: list[str], conf: Dict[str, Any]) -> None:
    style = conf["styles"]["code"]
    for line in lines:
        p = doc.add_paragraph()
        apply_paragraph_fmt(p, style)
        run = p.add_run(line if line.strip() else " ")
        apply_run_fmt(run, style, conf["global"])


def add_horizontal_rule(doc: Document, conf: Dict[str, Any]) -> None:
    """Add a horizontal rule (line divider) to the document."""
    mode = conf.get("global", {}).get("horizontalRule", "default")

    if mode == "hidden":
        return

    if mode == "page_break":
        doc.add_page_break()
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


DEFAULT_MARGIN = 1.0
DEFAULT_LR_MARGIN = 1.25

def set_page_margins(doc: Document, margin: float | Dict[str, float]) -> None:
    for section in doc.sections:
        if isinstance(margin, dict):
            section.top_margin = Inches(float(margin.get("top", DEFAULT_MARGIN)))
            section.bottom_margin = Inches(float(margin.get("bottom", DEFAULT_MARGIN)))
            section.left_margin = Inches(float(margin.get("left", DEFAULT_LR_MARGIN)))
            section.right_margin = Inches(float(margin.get("right", DEFAULT_LR_MARGIN)))
        else:
            m_val = float(margin)
            section.top_margin = Inches(m_val)
            section.bottom_margin = Inches(m_val)
            section.left_margin = Inches(m_val)
            section.right_margin = Inches(m_val)


# ---------------------------------------------------------------------------
# Backward-compatible re-exports from converters
# ---------------------------------------------------------------------------
from .converters.table import add_table  # noqa: F401
from .converters.toc import add_table_of_contents  # noqa: F401
