"""Table conversion — detection, parsing, and DOCX rendering.

Centralises all table-specific logic that was previously scattered across
parser.py, elements.py, and converter.py.  Downstream code should import
from here for table operations; the original modules re-export for backward
compatibility.
"""

from typing import Any, Dict, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..parser import (
    parse_gfm_table,
    is_table_line,
    is_table_separator,
)
from ..converters.styles import (
    apply_paragraph_fmt, _get_alignment, _ensure_east_asia_font,
)
from ..parser import parse_inline_formatting


# ---------------------------------------------------------------------------
# Cell helpers
# ---------------------------------------------------------------------------

def _set_cell_border(cell, border_color: str = "000000", border_size: int = 4) -> None:
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, fill_color: str) -> None:
    """Set background shading for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_color)
    tcPr.append(shading)


def _set_cell_vertical_alignment(cell, align: str = "center") -> None:
    """Set vertical alignment for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


# ---------------------------------------------------------------------------
# Inline formatting helpers (local to table rendering)
# ---------------------------------------------------------------------------

def _add_hyperlink(paragraph, url: str, text: str, style_config: Dict[str, Any], global_config: Dict[str, Any]):
    """Add a hyperlink to a paragraph."""
    from docx.shared import RGBColor
    from ..converters.styles import apply_run_fmt

    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    run = paragraph.add_run(text)
    apply_run_fmt(run, style_config, global_config)

    run.font.color.rgb = RGBColor(0, 0, 255)
    run.underline = True

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)


def _add_formatted_runs(paragraph, text: str, base_style: Dict[str, Any], global_config: Dict[str, Any], code_style: Dict[str, Any] = None) -> None:
    """Add text with inline formatting to a paragraph (table-local copy)."""
    from ..converters.styles import apply_run_fmt, _set_run_shading

    segments = parse_inline_formatting(text)

    for segment in segments:
        if segment['link']:
            _add_hyperlink(paragraph, segment['link'], segment['text'], base_style, global_config)
        else:
            run = paragraph.add_run(segment['text'])

            if segment['code'] and code_style:
                apply_run_fmt(run, code_style, global_config)
            else:
                apply_run_fmt(run, base_style, global_config)

            if segment.get('bold'):
                run.bold = True
            if segment.get('italic'):
                run.italic = True
            if segment.get('underline'):
                run.underline = True
            if segment.get('strike'):
                run.font.strike = True
            if segment.get('code'):
                code_font = code_style.get('fontFamily', 'Courier New') if code_style else 'Courier New'
                run.font.name = code_font
                _ensure_east_asia_font(run, global_config.get('baseFontCn', 'SimSun'), code_font)
                if code_style:
                    bg_color = code_style.get('backgroundColor', '#F5F7F9')
                    _set_run_shading(run, bg_color)


# ---------------------------------------------------------------------------
# Public: add_table
# ---------------------------------------------------------------------------

def add_table(doc: Document, table_data: list[list[str]], conf: Dict[str, Any], alignments: Optional[list[Optional[str]]] = None) -> None:
    """Add a table to the document with proper borders and header styling."""
    if not table_data or not table_data[0]:
        return

    num_rows = len(table_data)
    num_cols = max(len(row) for row in table_data)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    style = conf["styles"].get("body", {})
    global_conf = conf.get("global", {})

    for row_idx, row_data in enumerate(table_data):
        row = table.rows[row_idx]
        for col_idx in range(num_cols):
            cell = row.cells[col_idx]
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""

            _set_cell_border(cell)
            _set_cell_vertical_alignment(cell, "center")

            if row_idx == 0:
                _set_cell_shading(cell, "E5E7EB")

            paragraph = cell.paragraphs[0]

            table_style = style.copy()
            table_style["firstLineIndent"] = 0
            apply_paragraph_fmt(paragraph, table_style)

            code_style = conf["styles"].get("code", {})
            _add_formatted_runs(paragraph, cell_text, table_style, global_conf, code_style)

            if alignments and col_idx < len(alignments):
                align = _get_alignment(alignments[col_idx])
                if align is not None:
                    paragraph.alignment = align

            if row_idx == 0:
                for run in paragraph.runs:
                    run.bold = True


# ---------------------------------------------------------------------------
# Buffered processing — the table buffering loop previously in converter.py
# ---------------------------------------------------------------------------

def flush_table_buffer(
    doc: Document,
    table_buf: list[str],
    conf: Dict[str, Any],
) -> None:
    """Parse accumulated table lines and add the table to *doc*.

    Does nothing if *table_buf* is empty or the content is not a valid GFM
    table.
    """
    if not table_buf:
        return
    table_data = parse_gfm_table(table_buf)
    if table_data:
        add_table(doc, table_data["rows"], conf, table_data["alignments"])


def process_table_buffer(
    doc: Document,
    line: str,
    in_table: bool,
    table_buf: list[str],
    conf: Dict[str, Any],
    next_line: Optional[str] = None,
) -> tuple[bool, list[str], bool]:
    """Run one iteration of the table-buffer state machine.

    Parameters
    ----------
    doc : Document
        The Word document being built.
    line : str
        Current line (already ``rstrip("\\n")``).
    in_table : bool
        Whether we are currently inside a table.
    table_buf : list[str]
        Lines accumulated so far for the current table.
    conf : dict
        Style configuration.
    next_line : str or None
        The line following *line* (used for table-start lookahead).

    Returns
    -------
    (consumed, table_buf, in_table) : tuple
        *consumed* is ``True`` when *line* was handled (caller should skip
        further processing and advance).  *table_buf* and *in_table* are the
        updated state.
    """
    # --- detect table start ---
    if not in_table:
        if is_table_line(line) and next_line is not None and is_table_separator(next_line):
            return True, [line], True
        return False, table_buf, in_table

    # --- inside a table: accumulate or flush ---
    if is_table_line(line) or is_table_separator(line):
        table_buf.append(line)
        return True, table_buf, in_table

    # line is not a table row → flush and fall through
    flush_table_buffer(doc, table_buf, conf)
    return False, [], False


# ---------------------------------------------------------------------------
# Re-exports for backward compatibility
# ---------------------------------------------------------------------------
__all__ = [
    "parse_gfm_table",
    "is_table_line",
    "is_table_separator",
    "add_table",
    "flush_table_buffer",
    "process_table_buffer",
]
