"""Table of Contents conversion — detection and DOCX rendering.

Centralises TOC-specific logic so that converter.py does not need to
know the internals of TOC field injection.
"""

import sys
from typing import Any, Dict

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from ..converters.styles import _ensure_east_asia_font, _get_alignment, hex_to_rgb


def _set_style_fonts(style, cn_font: str, en_font: str) -> None:
    style.font.name = en_font
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)


def _apply_style_to_run(run, style_config: Dict[str, Any], global_conf: Dict[str, Any]) -> None:
    font_size = style_config.get("fontSize")
    if font_size is not None:
        run.font.size = Pt(float(font_size))
    if style_config.get("bold") is not None:
        run.bold = bool(style_config.get("bold"))
    if style_config.get("italic") is not None:
        run.italic = bool(style_config.get("italic"))
    if style_config.get("color"):
        run.font.color.rgb = hex_to_rgb(str(style_config.get("color")))

    cn_font = style_config.get("fontFamily") or global_conf.get("baseFontCn", "SimSun")
    en_font = style_config.get("fontFamilyEn") or global_conf.get("baseFontEn") or cn_font
    _ensure_east_asia_font(run, cn_font, en_font)


def _apply_toc_entry_styles(doc: Document, conf: Dict[str, Any]) -> None:
    global_conf = conf.get("global", {})
    toc_conf = global_conf.get("tableOfContents") or {}
    level_styles = toc_conf.get("levelStyles") or {}

    for level, style_config in level_styles.items():
        style_name = f"TOC {int(level)}"
        try:
            style = doc.styles[style_name]
        except Exception:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        cn_font = style_config.get("fontFamily") or global_conf.get("baseFontCn", "SimSun")
        en_font = style_config.get("fontFamilyEn") or global_conf.get("baseFontEn") or cn_font
        _set_style_fonts(style, cn_font, en_font)

        if style_config.get("fontSize") is not None:
            style.font.size = Pt(float(style_config["fontSize"]))
        if style_config.get("bold") is not None:
            style.font.bold = bool(style_config["bold"])
        if style_config.get("italic") is not None:
            style.font.italic = bool(style_config["italic"])
        if style_config.get("color"):
            style.font.color.rgb = hex_to_rgb(str(style_config["color"]))

        paragraph_format = style.paragraph_format
        first_line_indent = style_config.get("firstLineIndent")
        if first_line_indent is not None:
            paragraph_format.left_indent = Inches((float(first_line_indent) * 12.0) / 72.0)

        align = _get_alignment(style_config.get("alignment"))
        if align is not None:
            paragraph_format.alignment = align


def should_add_toc(conf: Dict[str, Any]) -> bool:
    """Return True when the configuration requests a Table of Contents."""
    return bool(conf.get("global", {}).get("includeTableOfContents", False))


def add_table_of_contents(doc: Document, conf: Dict[str, Any]) -> None:
    """
    Add a Table of Contents field to the document.
    The TOC will be auto-updated when the document is opened in Word.
    """
    global_conf = conf.get("global", {})
    toc_conf = global_conf.get("tableOfContents") or {}
    title_style = toc_conf.get("titleStyle") or {}
    _apply_toc_entry_styles(doc, conf)

    toc_title = doc.add_paragraph()
    toc_title.paragraph_format.space_before = Pt(float(title_style.get("spaceBefore", 12)))
    toc_title.paragraph_format.space_after = Pt(float(title_style.get("spaceAfter", 12)))
    align = _get_alignment(title_style.get("alignment", "center"))
    toc_title.alignment = align if align is not None else WD_PARAGRAPH_ALIGNMENT.CENTER
    run = toc_title.add_run("目录")
    _apply_style_to_run(
        run,
        {
            "fontFamily": title_style.get("fontFamily", global_conf.get("baseFontCn", "SimSun")),
            "fontFamilyEn": title_style.get("fontFamilyEn", global_conf.get("baseFontEn")),
            "fontSize": title_style.get("fontSize", 16),
            "color": title_style.get("color", "#000000"),
            "bold": title_style.get("bold", True),
            "italic": title_style.get("italic", False),
        },
        global_conf,
    )

    toc_paragraph = doc.add_paragraph()
    max_level = int(toc_conf.get("maxLevel", 3) or 3)

    run = toc_paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = f' TOC \\o "1-{max_level}" \\h \\z \\u '

    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)

    placeholder_run = toc_paragraph.add_run('右键点击并选择"更新域"，或按 Ctrl+A 然后 F9')
    placeholder_run.italic = True
    placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
    placeholder_run.font.size = Pt(10)

    end_run = toc_paragraph.add_run()
    end_run._r.append(fld_char_end)

    body_start = global_conf.get("bodyStart") or {}
    if body_start.get("restartPageNumberAfterToc"):
        doc.add_section(WD_SECTION.NEW_PAGE)
    else:
        doc.add_page_break()


def add_toc(doc: Document, conf: Dict[str, Any]) -> None:
    """Add a Table of Contents to *doc* if the config requests one.

    Errors are printed to stderr but do **not** abort the conversion —
    matching the previous behaviour in converter.py.
    """
    if not should_add_toc(conf):
        return
    try:
        add_table_of_contents(doc, conf)
    except Exception as e:
        print(f"Warning: Failed to add table of contents: {e}", file=sys.stderr)


__all__ = ["should_add_toc", "add_toc", "add_table_of_contents"]
