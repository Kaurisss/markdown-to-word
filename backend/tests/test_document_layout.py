import os
import sys

from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from backend import convert


def _course_design_conf():
    return {
        "global": {
            "pageMargin": {
                "top": 3.5 / 2.54,
                "bottom": 3.0 / 2.54,
                "left": 3.0 / 2.54,
                "right": 2.5 / 2.54,
            },
            "pageSize": {"width": 21.0, "height": 29.7, "unit": "cm"},
            "baseFontCn": "SimSun",
            "baseFontEn": "Times New Roman",
            "horizontalRule": "default",
            "includeTableOfContents": True,
            "header": {
                "enabled": True,
                "text": "",
                "distance": 2.8 / 2.54,
                "fontFamily": "SimSun",
                "fontFamilyEn": "Times New Roman",
                "fontSize": 10.5,
                "alignment": "center",
            },
            "footer": {
                "enabled": True,
                "pageNumber": True,
                "format": "第{page}页（共{pages}页）",
                "distance": 2.2 / 2.54,
                "fontFamily": "SimSun",
                "fontFamilyEn": "Times New Roman",
                "fontSize": 10.5,
                "alignment": "center",
                "startAtBody": True,
            },
            "tableOfContents": {
                "maxLevel": 2,
                "titleStyle": {"fontFamily": "SimHei", "fontFamilyEn": "Times New Roman", "fontSize": 18, "color": "#000000", "bold": True, "italic": False, "spaceBefore": 12, "spaceAfter": 12, "alignment": "center"},
                "levelStyles": {
                    "1": {"fontFamily": "SimHei", "fontFamilyEn": "Times New Roman", "fontSize": 12, "color": "#000000", "bold": True, "italic": False, "alignment": "left", "firstLineIndent": 0},
                    "2": {"fontFamily": "SimSun", "fontFamilyEn": "Times New Roman", "fontSize": 12, "color": "#000000", "bold": False, "italic": False, "alignment": "left", "firstLineIndent": 2},
                },
            },
            "bodyStart": {"firstHeadingAsTitle": True, "restartPageNumberAfterToc": True, "pageNumberStart": 1},
            "tableHeaderBold": False,
            "normalizePunctuation": True,
        },
        "styles": {
            "documentTitle": {"fontFamily": "SimHei", "fontFamilyEn": "Times New Roman", "fontSize": 18, "color": "#000000", "bold": True, "italic": False, "lineSpacing": 1.2, "spaceBefore": 12, "spaceAfter": 12, "alignment": "center", "firstLineIndent": 0},
            "h1": {"fontFamily": "SimHei", "fontFamilyEn": "Times New Roman", "fontSize": 18, "color": "#000000", "bold": True, "italic": False, "lineSpacing": 1.2, "spaceBefore": 6, "spaceAfter": 6, "alignment": "left", "firstLineIndent": 0},
            "h2": {"fontFamily": "SimHei", "fontFamilyEn": "Times New Roman", "fontSize": 16, "color": "#000000", "bold": True, "italic": False, "lineSpacing": 1.2, "spaceBefore": 6, "spaceAfter": 6, "alignment": "left", "firstLineIndent": 0},
            "h3": {"fontFamily": "SimHei", "fontFamilyEn": "Times New Roman", "fontSize": 12, "color": "#000000", "bold": True, "italic": False, "lineSpacing": 1.2, "spaceBefore": 6, "spaceAfter": 6, "alignment": "left", "firstLineIndent": 0},
            "body": {"fontFamily": "SimSun", "fontFamilyEn": "Times New Roman", "fontSize": 12, "color": "#000000", "bold": False, "italic": False, "lineSpacing": "22pt", "spaceBefore": 0, "spaceAfter": 0, "alignment": "left", "firstLineIndent": 2},
            "code": {"fontFamily": "Courier New", "fontFamilyEn": "Courier New", "fontSize": 10, "color": "#000000", "bold": False, "italic": False, "lineSpacing": 1.2, "spaceBefore": 0, "spaceAfter": 0, "alignment": "left", "firstLineIndent": 0},
            "quote": {"fontFamily": "SimSun", "fontFamilyEn": "Times New Roman", "fontSize": 12, "color": "#000000", "bold": False, "italic": True, "lineSpacing": 1.4, "spaceBefore": 8, "spaceAfter": 8, "alignment": "left", "firstLineIndent": 0},
            "table": {"fontFamily": "SimSun", "fontFamilyEn": "Times New Roman", "fontSize": 10.5, "color": "#000000", "bold": False, "italic": False, "lineSpacing": 1.2, "spaceBefore": 0, "spaceAfter": 0, "alignment": "center", "firstLineIndent": 0},
            "caption": {"fontFamily": "KaiTi", "fontFamilyEn": "Times New Roman", "fontSize": 10.5, "color": "#000000", "bold": True, "italic": False, "lineSpacing": 1.2, "spaceBefore": 6, "spaceAfter": 6, "alignment": "center", "firstLineIndent": 0},
        },
    }


def _convert(tmp_path, markdown: str):
    input_path = tmp_path / "input.md"
    output_path = tmp_path / "output.docx"
    input_path.write_text(markdown, encoding="utf-8")
    convert(str(input_path), str(output_path), _course_design_conf())
    return Document(str(output_path))


def _rfonts(run):
    r_fonts = run._r.rPr.rFonts
    return {
        "eastAsia": r_fonts.get(qn("w:eastAsia")),
        "ascii": r_fonts.get(qn("w:ascii")),
        "hAnsi": r_fonts.get(qn("w:hAnsi")),
    }


def test_course_design_page_layout_header_footer_and_page_restart(tmp_path):
    doc = _convert(
        tmp_path,
        "\n".join([
            "# 课程设计题目",
            "",
            "## 第一章",
            "",
            "这是正文 ABC 123。",
        ]),
    )

    assert len(doc.sections) == 2

    first_section = doc.sections[0]
    body_section = doc.sections[1]

    assert abs(first_section.page_width.cm - 21.0) < 0.05
    assert abs(first_section.page_height.cm - 29.7) < 0.05
    assert abs(body_section.top_margin.cm - 3.5) < 0.05
    assert abs(body_section.bottom_margin.cm - 3.0) < 0.05
    assert abs(body_section.left_margin.cm - 3.0) < 0.05
    assert abs(body_section.right_margin.cm - 2.5) < 0.05

    assert first_section.header.paragraphs[0].text == ""
    assert abs(first_section.header_distance.cm - 2.8) < 0.05
    assert abs(body_section.footer_distance.cm - 2.2) < 0.05

    footer_xml = body_section.footer._element.xml
    assert "PAGE" in footer_xml
    assert "NUMPAGES" in footer_xml
    assert "第" in body_section.footer.paragraphs[0].text
    assert "页（共" in body_section.footer.paragraphs[0].text

    pg_num_type = body_section._sectPr.find(qn("w:pgNumType"))
    assert pg_num_type is not None
    assert pg_num_type.get(qn("w:start")) == "1"


def test_course_design_fonts_and_title_styles(tmp_path):
    doc = _convert(tmp_path, "# 课程设计题目\n\n## 第一章\n\n这是正文, ABC 123.")

    title = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "课程设计题目")
    assert title.style.name == "Title"
    assert title.alignment == 1
    assert title.runs[0].font.size.pt == 18
    title_fonts = _rfonts(title.runs[0])
    assert title_fonts["eastAsia"] == "SimHei"
    assert title_fonts["ascii"] == "Times New Roman"

    body = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "这是正文， ABC 123。")
    assert body.paragraph_format.line_spacing.pt == 22
    assert round(body.paragraph_format.first_line_indent.inches, 2) == 0.33

    body_fonts = _rfonts(body.runs[0])
    assert body_fonts["eastAsia"] == "SimSun"
    assert body_fonts["ascii"] == "Times New Roman"


def test_course_design_toc_and_table_text_styles(tmp_path):
    doc = _convert(
        tmp_path,
        "\n".join([
            "# 课程设计题目",
            "",
            "## 第一章",
            "",
            "| 中文 | Value |",
            "| --- | --- |",
            "| 内容 | ABC123 |",
            "",
            "图 示例图",
        ]),
    )

    assert doc.styles["TOC 1"].font.size.pt == 12
    toc_1_fonts = doc.styles["TOC 1"].element.rPr.rFonts
    assert toc_1_fonts.get(qn("w:eastAsia")) == "SimHei"
    assert doc.styles["TOC 2"].paragraph_format.left_indent.pt == 24

    header_run = doc.tables[0].cell(0, 0).paragraphs[0].runs[0]
    assert header_run.font.size.pt == 10.5
    assert header_run.bold is False
    table_fonts = _rfonts(header_run)
    assert table_fonts["eastAsia"] == "SimSun"
    assert table_fonts["ascii"] == "Times New Roman"

    caption = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "图1 示例图")
    caption_run = caption.runs[0]
    assert caption_run.font.size.pt == 10.5
    assert caption_run.bold is True
    caption_fonts = _rfonts(caption_run)
    assert caption_fonts["eastAsia"] == "KaiTi"
