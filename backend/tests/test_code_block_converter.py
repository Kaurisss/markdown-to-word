"""Unit tests for converters.code_block — process_code_buffer state machine.

Covers: fence detection, accumulation, flush on closing fence,
fallthrough when not in code, and flush at EOF.
"""
import os
import sys

from docx import Document

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from backend.converters.code_block import (
    flush_code_buffer, is_fence, process_code_buffer,
)


def minimal_conf():
    return {
        "global": {"pageMargin": 1.0},
        "styles": {
            "h1": {}, "h2": {}, "h3": {},
            "body": {}, "code": {}, "quote": {},
        },
    }


# ---------------------------------------------------------------------------
# is_fence
# ---------------------------------------------------------------------------

class TestIsFence:
    def test_detects_opening_fence(self):
        assert is_fence("```python") is True

    def test_detects_bare_fence(self):
        assert is_fence("```") is True

    def test_rejects_normal_text(self):
        assert is_fence("hello") is False

    def test_rejects_indented_fence(self):
        assert is_fence("  ```") is False


# ---------------------------------------------------------------------------
# process_code_buffer
# ---------------------------------------------------------------------------

class TestCodeStart:
    def test_detects_opening_fence(self):
        doc = Document()
        consumed, buf, in_code = process_code_buffer(
            doc, "```python", False, [], minimal_conf(),
        )
        assert consumed is True
        assert in_code is True
        assert buf == []

    def test_plain_text_falls_through(self):
        doc = Document()
        consumed, buf, in_code = process_code_buffer(
            doc, "hello", False, [], minimal_conf(),
        )
        assert consumed is False
        assert in_code is False


class TestCodeAccumulation:
    def test_accumulates_code_line(self):
        doc = Document()
        consumed, buf, in_code = process_code_buffer(
            doc, "print('hi')", True, ["x = 1"], minimal_conf(),
        )
        assert consumed is True
        assert in_code is True
        assert buf == ["x = 1", "print('hi')"]


class TestCodeFlush:
    def test_closing_fence_flushes(self):
        doc = Document()
        consumed, buf, in_code = process_code_buffer(
            doc, "```", True, ["x = 1", "print(x)"], minimal_conf(),
        )
        assert consumed is True
        assert in_code is False
        assert buf == []
        assert len(doc.paragraphs) >= 2

    def test_flush_empty_buffer_is_noop(self):
        doc = Document()
        flush_code_buffer(doc, [], minimal_conf())
        assert len(doc.paragraphs) == 0

    def test_flush_renders_code_lines(self):
        doc = Document()
        flush_code_buffer(doc, ["line1", "line2"], minimal_conf())
        texts = [p.text for p in doc.paragraphs]
        assert "line1" in texts
        assert "line2" in texts


class TestCodeFallthrough:
    def test_not_in_code_falls_through(self):
        doc = Document()
        consumed, buf, in_code = process_code_buffer(
            doc, "just text", False, [], minimal_conf(),
        )
        assert consumed is False
        assert in_code is False
        assert buf == []
