"""
Property-based tests for inline formatting.

**Feature: python-backend-upgrade, Property 9: Inline Formatting Application**
**Validates: Requirements 3.7**
"""
import os
import sys
import tempfile

from hypothesis import given, strategies as st, settings, assume
from docx import Document

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from backend import parse_inline_formatting, convert, load_config


# Strategy for generating plain text (no special markdown characters)
plain_text = st.text(
    alphabet=st.characters(
        whitelist_categories=('L', 'N', 'S'),
        blacklist_characters='*`[]()|\n\r'
    ),
    min_size=1,
    max_size=20
).filter(lambda x: x.strip() != '')


@settings(max_examples=100)
@given(text=plain_text)
def test_bold_formatting_preserved(text):
    """
    **Feature: python-backend-upgrade, Property 9: Inline Formatting Application**
    **Validates: Requirements 3.7**
    
    For any text wrapped in **bold** markers, the parser SHALL identify it as bold.
    """
    markdown = f"**{text}**"
    segments = parse_inline_formatting(markdown)
    
    # Should have exactly one segment
    assert len(segments) == 1, f"Expected 1 segment, got {len(segments)}"
    
    # The segment should be bold
    assert segments[0]['bold'] is True, "Text should be marked as bold"
    assert segments[0]['text'] == text, f"Text content mismatch: expected '{text}', got '{segments[0]['text']}'"


@settings(max_examples=100)
@given(text=plain_text)
def test_italic_formatting_preserved(text):
    """
    **Feature: python-backend-upgrade, Property 9: Inline Formatting Application**
    **Validates: Requirements 3.7**
    
    For any text wrapped in *italic* markers, the parser SHALL identify it as italic.
    """
    markdown = f"*{text}*"
    segments = parse_inline_formatting(markdown)
    
    # Should have exactly one segment
    assert len(segments) == 1, f"Expected 1 segment, got {len(segments)}"
    
    # The segment should be italic
    assert segments[0]['italic'] is True, "Text should be marked as italic"
    assert segments[0]['text'] == text, f"Text content mismatch: expected '{text}', got '{segments[0]['text']}'"


@settings(max_examples=100)
@given(text=plain_text)
def test_code_formatting_preserved(text):
    """
    **Feature: python-backend-upgrade, Property 9: Inline Formatting Application**
    **Validates: Requirements 3.7**
    
    For any text wrapped in `code` markers, the parser SHALL identify it as code.
    """
    markdown = f"`{text}`"
    segments = parse_inline_formatting(markdown)
    
    # Should have exactly one segment
    assert len(segments) == 1, f"Expected 1 segment, got {len(segments)}"
    
    # The segment should be code
    assert segments[0]['code'] is True, "Text should be marked as code"
    assert segments[0]['text'] == text, f"Text content mismatch: expected '{text}', got '{segments[0]['text']}'"


# Strategy for generating valid URLs
url_strategy = st.from_regex(r'https?://[a-z]+\.[a-z]+', fullmatch=True)


@settings(max_examples=100)
@given(text=plain_text, url=url_strategy)
def test_link_formatting_preserved(text, url):
    """
    **Feature: python-backend-upgrade, Property 9: Inline Formatting Application**
    **Validates: Requirements 3.7**
    
    For any text in [text](url) format, the parser SHALL identify it as a link.
    """
    markdown = f"[{text}]({url})"
    segments = parse_inline_formatting(markdown)
    
    # Should have exactly one segment
    assert len(segments) == 1, f"Expected 1 segment, got {len(segments)}"
    
    # The segment should be a link
    assert segments[0]['link'] == url, f"URL mismatch: expected '{url}', got '{segments[0]['link']}'"
    assert segments[0]['text'] == text, f"Text content mismatch: expected '{text}', got '{segments[0]['text']}'"


@settings(max_examples=50)
@given(
    prefix=plain_text,
    bold_text=plain_text,
    middle=plain_text,
    italic_text=plain_text,
    suffix=plain_text
)
def test_mixed_formatting_preserved(prefix, bold_text, middle, italic_text, suffix):
    """
    **Feature: python-backend-upgrade, Property 9: Inline Formatting Application**
    **Validates: Requirements 3.7**
    
    For any text with mixed inline formatting, the parser SHALL correctly identify
    all formatting types and preserve text content.
    """
    markdown = f"{prefix} **{bold_text}** {middle} *{italic_text}* {suffix}"
    segments = parse_inline_formatting(markdown)
    
    # Find bold and italic segments
    bold_segments = [s for s in segments if s['bold']]
    italic_segments = [s for s in segments if s['italic']]
    
    # Should have at least one bold and one italic segment
    assert len(bold_segments) >= 1, "Should have at least one bold segment"
    assert len(italic_segments) >= 1, "Should have at least one italic segment"
    
    # Verify bold text content
    assert any(s['text'] == bold_text for s in bold_segments), \
        f"Bold text '{bold_text}' not found in segments"
    
    # Verify italic text content
    assert any(s['text'] == italic_text for s in italic_segments), \
        f"Italic text '{italic_text}' not found in segments"


@settings(max_examples=30)
@given(bold_text=plain_text)
def test_inline_formatting_renders_to_docx(bold_text):
    """
    **Feature: python-backend-upgrade, Property 9: Inline Formatting Application**
    **Validates: Requirements 3.7**
    
    For any text with inline formatting, the generated Word document SHALL
    contain runs with the corresponding formatting applied.
    """
    markdown_content = f"This is **{bold_text}** text."
    
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, 'input.md')
        output_path = os.path.join(tmpdir, 'output.docx')
        
        # Write markdown file
        with open(input_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        # Create a mock args object for load_config
        class Args:
            config = None
            config_file = None
        
        conf = load_config(Args())
        
        # Convert to docx
        convert(input_path, output_path, conf)
        
        # Verify output file exists
        assert os.path.exists(output_path), "Output file should be created"
        
        # Open and verify the document
        doc = Document(output_path)
        
        # Should have at least one paragraph
        assert len(doc.paragraphs) >= 1, "Document should have at least one paragraph"
        
        # Find the paragraph with our content
        found_bold = False
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text == bold_text and run.bold:
                    found_bold = True
                    break
        
        assert found_bold, f"Bold text '{bold_text}' not found with bold formatting in document"
