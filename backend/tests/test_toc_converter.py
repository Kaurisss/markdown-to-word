"""Unit tests for converters.toc — should_add_toc and add_toc."""
import os
import sys

from docx import Document

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from backend.converters.toc import should_add_toc, add_toc


def conf_with_toc(enabled=True):
    return {
        "global": {"pageMargin": 1.0, "includeTableOfContents": enabled},
        "styles": {
            "h1": {}, "h2": {}, "h3": {},
            "body": {}, "code": {}, "quote": {},
        },
    }


def conf_without_toc_key():
    return {
        "global": {"pageMargin": 1.0},
        "styles": {
            "h1": {}, "h2": {}, "h3": {},
            "body": {}, "code": {}, "quote": {},
        },
    }


# ---------------------------------------------------------------------------
# should_add_toc
# ---------------------------------------------------------------------------

class TestShouldAddToc:
    def test_true_when_enabled(self):
        assert should_add_toc(conf_with_toc(True)) is True

    def test_false_when_disabled(self):
        assert should_add_toc(conf_with_toc(False)) is False

    def test_false_when_key_missing(self):
        assert should_add_toc(conf_without_toc_key()) is False


# ---------------------------------------------------------------------------
# add_toc
# ---------------------------------------------------------------------------

class TestAddToc:
    def test_adds_toc_when_enabled(self):
        doc = Document()
        add_toc(doc, conf_with_toc(True))
        # TOC inserts a page break, so there should be at least 2 paragraphs
        # (title + TOC field paragraph + page break paragraph)
        assert len(doc.paragraphs) >= 2

    def test_noop_when_disabled(self):
        doc = Document()
        initial = len(doc.paragraphs)
        add_toc(doc, conf_with_toc(False))
        assert len(doc.paragraphs) == initial

    def test_noop_when_key_missing(self):
        doc = Document()
        initial = len(doc.paragraphs)
        add_toc(doc, conf_without_toc_key())
        assert len(doc.paragraphs) == initial

    def test_toc_title_text(self):
        doc = Document()
        add_toc(doc, conf_with_toc(True))
        # First paragraph should be the TOC title
        assert "目录" in doc.paragraphs[0].text

    def test_toc_does_not_abort_on_minimal_config(self):
        """Even with a minimal config (no styles), add_toc should not raise."""
        doc = Document()
        minimal_conf = {"global": {"includeTableOfContents": True}}
        add_toc(doc, minimal_conf)
        assert len(doc.paragraphs) >= 2
