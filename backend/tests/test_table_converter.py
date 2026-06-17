"""Unit tests for converters.table.process_table_buffer state machine.

Covers: start detection, accumulation, flush on non-table line,
flush at EOF, and fallthrough when not in table.
"""
import os
import sys

import pytest
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from backend.converters.table import process_table_buffer, flush_table_buffer


def minimal_conf():
    return {
        "global": {"pageMargin": 1.0},
        "styles": {
            "h1": {}, "h2": {}, "h3": {},
            "body": {}, "code": {}, "quote": {},
        },
    }


# ---------------------------------------------------------------------------
# process_table_buffer
# ---------------------------------------------------------------------------

class TestTableStart:
    """Detect table start when current line looks like a table row and the
    next line is a separator."""

    def test_detects_table_start(self):
        doc = Document()
        consumed, buf, in_table = process_table_buffer(
            doc, "| A | B |", False, [], minimal_conf(), "| --- | --- |",
        )
        assert consumed is True
        assert in_table is True
        assert buf == ["| A | B |"]

    def test_no_start_without_next_line(self):
        doc = Document()
        consumed, buf, in_table = process_table_buffer(
            doc, "| A | B |", False, [], minimal_conf(), next_line=None,
        )
        assert consumed is False
        assert in_table is False

    def test_no_start_when_next_not_separator(self):
        doc = Document()
        consumed, buf, in_table = process_table_buffer(
            doc, "| A | B |", False, [], minimal_conf(), "just text",
        )
        assert consumed is False
        assert in_table is False


class TestTableAccumulation:
    """While in_table, table-like and separator lines are accumulated."""

    def test_accumulates_table_row(self):
        doc = Document()
        consumed, buf, in_table = process_table_buffer(
            doc, "| X | Y |", True, ["| A | B |"], minimal_conf(),
        )
        assert consumed is True
        assert in_table is True
        assert buf == ["| A | B |", "| X | Y |"]

    def test_accumulates_separator(self):
        doc = Document()
        consumed, buf, in_table = process_table_buffer(
            doc, "| --- | --- |", True, ["| A | B |"], minimal_conf(),
        )
        assert consumed is True
        assert in_table is True


class TestTableFlush:
    """When a non-table line arrives while in_table, the buffer is flushed
    and the line is NOT consumed (fallthrough)."""

    def test_flushes_and_falls_through(self):
        doc = Document()
        table_buf = ["| A | B |", "| --- | --- |", "| 1 | 2 |"]
        consumed, buf, in_table = process_table_buffer(
            doc, "some paragraph", True, table_buf, minimal_conf(),
        )
        assert consumed is False
        assert in_table is False
        assert buf == []
        # Document should now contain the table
        assert len(doc.tables) == 1

    def test_flush_empty_buffer_is_noop(self):
        doc = Document()
        consumed, buf, in_table = process_table_buffer(
            doc, "text", True, [], minimal_conf(),
        )
        assert consumed is False
        assert in_table is False
        assert len(doc.tables) == 0


class TestTableFallthrough:
    """When not in_table and line is not a table start, nothing happens."""

    def test_plain_text_falls_through(self):
        doc = Document()
        consumed, buf, in_table = process_table_buffer(
            doc, "hello world", False, [], minimal_conf(),
        )
        assert consumed is False
        assert in_table is False
        assert buf == []


class TestFlushTableBuffer:
    """Direct tests for flush_table_buffer."""

    def test_flush_valid_table(self):
        doc = Document()
        table_buf = ["| A | B |", "| --- | --- |", "| 1 | 2 |"]
        flush_table_buffer(doc, table_buf, minimal_conf())
        assert len(doc.tables) == 1
        assert len(doc.tables[0].rows) == 2

    def test_flush_empty_buffer(self):
        doc = Document()
        flush_table_buffer(doc, [], minimal_conf())
        assert len(doc.tables) == 0

    def test_flush_invalid_lines(self):
        doc = Document()
        flush_table_buffer(doc, ["not a table"], minimal_conf())
        assert len(doc.tables) == 0


class TestTableBeforeCodeFence:
    """Regression: a table immediately followed by a code fence must render
    the table BEFORE the code paragraphs in the output document."""

    def test_table_before_code_block(self, tmp_path):
        from backend import convert, load_config

        md = (
            "| A | B |\n"
            "| --- | --- |\n"
            "| 1 | 2 |\n"
            "```\n"
            "some code\n"
            "```\n"
        )
        inp = tmp_path / "input.md"
        outp = tmp_path / "output.docx"
        inp.write_text(md, encoding="utf-8")

        class Args:
            config = None
            config_file = None

        conf = load_config(Args())
        convert(str(inp), str(outp), conf)

        from docx import Document as DocxDocument
        doc = DocxDocument(str(outp))

        # Table must exist
        assert len(doc.tables) == 1, "Expected exactly 1 table"

        # In python-docx, tables and paragraphs are interleaved in
        # document order.  The table element comes first in the XML
        # body when it was added before the code paragraphs.
        body = doc.element.body
        children = list(body)
        # Find first tbl and first w:p that contains "some code"
        tbl_idx = next(
            (i for i, c in enumerate(children) if c.tag.endswith("}tbl")),
            None,
        )
        code_idx = next(
            (i for i, c in enumerate(children)
             if c.tag.endswith("}p") and c.text and "some code" in c.text),
            None,
        )
        assert tbl_idx is not None, "Table not found in document body"
        assert code_idx is not None, "Code paragraph not found"
        assert tbl_idx < code_idx, (
            f"Table (idx {tbl_idx}) should appear before code (idx {code_idx})"
        )
