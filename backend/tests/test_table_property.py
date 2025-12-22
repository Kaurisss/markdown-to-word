"""
Property-based tests for table rendering.

**Feature: python-backend-upgrade, Property 8: Table Structure Preservation**
**Validates: Requirements 3.6**
"""
import os
import sys
import tempfile

from hypothesis import given, strategies as st, settings
from docx import Document

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from backend import parse_gfm_table, convert, load_config


# Strategy for generating valid table cell content (non-empty, no pipes or newlines)
cell_content = st.text(
    alphabet=st.characters(
        whitelist_categories=('L', 'N', 'P', 'S'),
        blacklist_characters='|\n\r'
    ),
    min_size=1,
    max_size=20
).filter(lambda x: x.strip() != '')


# Strategy for generating table data (2D list of strings)
@st.composite
def table_data_strategy(draw):
    """Generate valid table data with consistent column counts."""
    num_cols = draw(st.integers(min_value=1, max_value=5))
    num_rows = draw(st.integers(min_value=1, max_value=10))
    
    rows = []
    for _ in range(num_rows):
        row = [draw(cell_content) for _ in range(num_cols)]
        rows.append(row)
    return rows


def table_data_to_gfm_lines(table_data: list[list[str]]) -> list[str]:
    """Convert table data to GFM markdown lines."""
    if not table_data:
        return []
    
    lines = []
    # Header row
    header = '| ' + ' | '.join(table_data[0]) + ' |'
    lines.append(header)
    
    # Separator row
    num_cols = len(table_data[0])
    separator = '| ' + ' | '.join(['---'] * num_cols) + ' |'
    lines.append(separator)
    
    # Data rows
    for row in table_data[1:]:
        line = '| ' + ' | '.join(row) + ' |'
        lines.append(line)
    
    return lines


@settings(max_examples=100)
@given(table_data=table_data_strategy())
def test_table_structure_preservation(table_data):
    """
    **Feature: python-backend-upgrade, Property 8: Table Structure Preservation**
    **Validates: Requirements 3.6**
    
    For any Markdown content containing GFM tables, the Python_Backend SHALL render
    tables with the correct number of rows and columns, proper cell borders, and
    header row styling.
    """
    # Convert table data to GFM lines
    gfm_lines = table_data_to_gfm_lines(table_data)
    
    # Parse the GFM table
    parsed = parse_gfm_table(gfm_lines)
    
    # Verify parsing preserves structure
    assert parsed is not None, "Table should be parsed successfully"
    assert len(parsed) == len(table_data), f"Row count mismatch: expected {len(table_data)}, got {len(parsed)}"
    
    for i, (expected_row, actual_row) in enumerate(zip(table_data, parsed)):
        assert len(actual_row) == len(expected_row), f"Column count mismatch in row {i}"
        for j, (expected_cell, actual_cell) in enumerate(zip(expected_row, actual_row)):
            assert actual_cell == expected_cell, f"Cell content mismatch at ({i}, {j})"


@settings(max_examples=50)
@given(table_data=table_data_strategy())
def test_table_renders_to_docx(table_data):
    """
    **Feature: python-backend-upgrade, Property 8: Table Structure Preservation**
    **Validates: Requirements 3.6**
    
    For any valid table data, the generated Word document SHALL contain a table
    with the correct number of rows and columns.
    """
    # Create markdown content with table
    gfm_lines = table_data_to_gfm_lines(table_data)
    markdown_content = '\n'.join(gfm_lines)
    
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, 'input.md')
        output_path = os.path.join(tmpdir, 'output.docx')
        
        # Write markdown file
        with open(input_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        # Create a mock args object for load_config
        class Args:
            config = None
            config_file = None
        
        conf = load_config(Args())
        
        # Convert to docx
        convert(input_path, output_path, conf)
        
        # Verify output file exists
        assert os.path.exists(output_path), "Output file should be created"
        
        # Open and verify the document
        doc = Document(output_path)
        
        # Should have exactly one table
        assert len(doc.tables) == 1, f"Expected 1 table, got {len(doc.tables)}"
        
        table = doc.tables[0]
        
        # Verify row count
        assert len(table.rows) == len(table_data), \
            f"Row count mismatch: expected {len(table_data)}, got {len(table.rows)}"
        
        # Verify column count
        expected_cols = len(table_data[0])
        for row_idx, row in enumerate(table.rows):
            assert len(row.cells) == expected_cols, \
                f"Column count mismatch in row {row_idx}: expected {expected_cols}, got {len(row.cells)}"
