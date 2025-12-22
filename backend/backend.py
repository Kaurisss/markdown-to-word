import argparse
import json
import os
import re
import sys
from typing import Any, Dict, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.table import Table


# Exit codes as defined in design document
EXIT_FILE_NOT_FOUND = 1
EXIT_PERMISSION_ERROR = 2
EXIT_CONFIG_ERROR = 3
EXIT_MARKDOWN_PARSE_ERROR = 4
EXIT_DOCX_GENERATION_ERROR = 5


class ConversionError(Exception):
    """Base exception for Markdown conversion errors."""
    exit_code = EXIT_MARKDOWN_PARSE_ERROR
    
    def __init__(self, message: str, details: str = None):
        self.message = message
        self.details = details
        super().__init__(message)
    
    def __str__(self):
        if self.details:
            return f"{self.message}: {self.details}"
        return self.message


class FileError(Exception):
    """Exception for file operation errors."""
    exit_code = EXIT_FILE_NOT_FOUND
    
    def __init__(self, message: str, path: str = None, details: str = None):
        self.message = message
        self.path = path
        self.details = details
        super().__init__(message)
    
    def __str__(self):
        parts = [self.message]
        if self.path:
            parts.append(f"Path: {self.path}")
        if self.details:
            parts.append(f"Details: {self.details}")
        return " - ".join(parts)


class ConfigError(Exception):
    """Exception for configuration parsing errors."""
    exit_code = EXIT_CONFIG_ERROR
    
    def __init__(self, message: str, details: str = None):
        self.message = message
        self.details = details
        super().__init__(message)
    
    def __str__(self):
        if self.details:
            return f"{self.message}: {self.details}"
        return self.message


class DocxGenerationError(Exception):
    """Exception for DOCX generation errors."""
    exit_code = EXIT_DOCX_GENERATION_ERROR
    
    def __init__(self, message: str, details: str = None):
        self.message = message
        self.details = details
        super().__init__(message)
    
    def __str__(self):
        if self.details:
            return f"{self.message}: {self.details}"
        return self.message


REQUIRED_STYLE_KEYS = ("h1", "h2", "h3", "body", "code", "quote")


def validate_config(conf: Dict[str, Any]) -> None:
    if not isinstance(conf, dict):
        raise ConfigError("Invalid configuration format", details="Expected a JSON object")

    global_conf = conf.get("global")
    styles = conf.get("styles")
    if not isinstance(global_conf, dict):
        raise ConfigError("Invalid configuration format", details="Missing or invalid 'global' section")
    if not isinstance(styles, dict):
        raise ConfigError("Invalid configuration format", details="Missing or invalid 'styles' section")

    missing = [key for key in REQUIRED_STYLE_KEYS if key not in styles]
    if missing:
        raise ConfigError("Invalid configuration format", details=f"Missing style keys: {', '.join(missing)}")

    for key in REQUIRED_STYLE_KEYS:
        if not isinstance(styles.get(key), dict):
            raise ConfigError("Invalid configuration format", details=f"Style '{key}' must be an object")

    page_margin = global_conf.get("pageMargin", 1.0)
    try:
        float(page_margin)
    except Exception as e:
        raise ConfigError("Invalid pageMargin value", details=str(e))


def _set_paragraph_shading(paragraph, color_hex: str):
    """Set background shading for a paragraph."""
    if not color_hex:
        return
    # Remove # if present
    fill_color = str(color_hex).strip().lstrip("#")
    
    pPr = paragraph._p.get_or_add_pPr()
    # Check if shd already exists to avoid duplication (though append usually works, safer to remove old)
    existing_shd = pPr.find(qn('w:shd'))
    if existing_shd is not None:
        pPr.remove(existing_shd)
        
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    pPr.append(shd)


def hex_to_rgb(hex_str: str) -> RGBColor:
    s = hex_str.strip().lstrip("#")
    if len(s) == 3:
        try:
            r = int(s[0] * 2, 16)
            g = int(s[1] * 2, 16)
            b = int(s[2] * 2, 16)
            return RGBColor(r, g, b)
        except Exception:
            return RGBColor(0, 0, 0)
    if len(s) != 6:
        return RGBColor(0, 0, 0)
    try:
        r = int(s[0:2], 16)
        g = int(s[2:4], 16)
        b = int(s[4:6], 16)
        return RGBColor(r, g, b)
    except Exception:
        return RGBColor(0, 0, 0)


def _get_alignment(value: Optional[str]) -> Optional[int]:
    if not value:
        return None
    m = {
        "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
    }
    return m.get(value.lower())


def apply_paragraph_fmt(paragraph, style_config: Dict[str, Any]) -> None:
    pf = paragraph.paragraph_format
    line_spacing = style_config.get("lineSpacing")
    if line_spacing is not None:
        try:
            # Handle fixed value (e.g., "20pt")
            if isinstance(line_spacing, str) and str(line_spacing).lower().endswith("pt"):
                val = float(str(line_spacing).lower().replace("pt", "").strip())
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(val)
            else:
                # Handle multiplier (number or string representation of number)
                pf.line_spacing = float(line_spacing)
        except Exception:
            pass
    space_before = style_config.get("spaceBefore")
    if space_before is not None:
        try:
            pf.space_before = Pt(float(space_before))
        except Exception:
            pass
    space_after = style_config.get("spaceAfter")
    if space_after is not None:
        try:
            pf.space_after = Pt(float(space_after))
        except Exception:
            pass
    align = _get_alignment(style_config.get("alignment"))
    if align is not None:
        pf.alignment = align
    first_indent_chars = style_config.get("firstLineIndent")
    if first_indent_chars is not None:
        try:
            # 约 12pt/字符 ≈ 12/72 英寸
            inches = (float(first_indent_chars) * 12.0) / 72.0
            pf.first_line_indent = Inches(inches)
        except Exception:
            pass
            
    bg_color = style_config.get("backgroundColor")
    if bg_color:
        _set_paragraph_shading(paragraph, bg_color)


def _ensure_east_asia_font(run, cn_font: str, en_font: str) -> None:
    run.font.name = en_font
    r = run._element
    rPr = r.rPr
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.append(rPr)
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)


def apply_run_fmt(run, style_config: Dict[str, Any], global_config: Dict[str, Any]) -> None:
    font_size = style_config.get("fontSize")
    if font_size:
        try:
            run.font.size = Pt(float(font_size))
        except Exception:
            pass
    if style_config.get("bold") is not None:
        run.bold = bool(style_config.get("bold"))
    if style_config.get("italic") is not None:
        run.italic = bool(style_config.get("italic"))
    color_hex = style_config.get("color")
    if color_hex:
        run.font.color.rgb = hex_to_rgb(str(color_hex))
    
    # 字体设置
    base_cn = (global_config.get("baseFontCn") or "SimSun")
    base_en = (global_config.get("baseFontEn") or base_cn)
    
    # 如果样式指定了 fontFamily，则覆盖中文字体
    ff = style_config.get("fontFamily")
    if ff:
        # 用户自定义字体同时用于中文和英文
        _ensure_east_asia_font(run, ff, base_en if base_en else ff)
    else:
        _ensure_east_asia_font(run, base_cn, base_en if base_en else base_cn)


def load_config(args) -> Dict[str, Any]:
    """Load configuration from file or JSON string with proper error handling."""
    if args.config_file:
        if not os.path.exists(args.config_file):
            raise FileError(
                "Configuration file not found",
                path=args.config_file
            )
        try:
            with open(args.config_file, "r", encoding="utf-8") as f:
                conf = json.load(f)
                validate_config(conf)
                return conf
        except json.JSONDecodeError as e:
            raise ConfigError(
                "Invalid JSON in configuration file",
                details=str(e)
            )
        except PermissionError as e:
            raise FileError(
                "Permission denied reading configuration file",
                path=args.config_file,
                details=str(e)
            )
    if args.config:
        try:
            conf = json.loads(args.config)
            validate_config(conf)
            return conf
        except json.JSONDecodeError as e:
            raise ConfigError(
                "Invalid JSON in configuration string",
                details=str(e)
            )
    # 默认中文配置
    conf = {
        "global": {
            "pageMargin": 1.0,
            "baseFontCn": "SimSun",
            "baseFontEn": "Times New Roman",
        },
        "styles": {
            "h1": {"fontSize": 24, "color": "#1F2937", "bold": True, "italic": False, "lineSpacing": 1.2, "spaceBefore": 12, "spaceAfter": 6, "alignment": "left", "firstLineIndent": 0},
            "h2": {"fontSize": 20, "color": "#1F2937", "bold": True, "italic": False, "lineSpacing": 1.2, "spaceBefore": 12, "spaceAfter": 6, "alignment": "left", "firstLineIndent": 0},
            "h3": {"fontSize": 18, "color": "#1F2937", "bold": True, "italic": False, "lineSpacing": 1.2, "spaceBefore": 12, "spaceAfter": 6, "alignment": "left", "firstLineIndent": 0},
            "body": {"fontSize": 12, "color": "#000000", "bold": False, "italic": False, "lineSpacing": 1.6, "spaceBefore": 0, "spaceAfter": 8, "alignment": "left", "firstLineIndent": 2},
            "code": {"fontSize": 10, "color": "#374151", "bold": False, "italic": False, "lineSpacing": 1.2, "spaceBefore": 0, "spaceAfter": 0, "alignment": "left", "firstLineIndent": 0, "fontFamily": "Courier New"},
            "quote": {"fontSize": 12, "color": "#4B5563", "bold": False, "italic": True, "lineSpacing": 1.4, "spaceBefore": 8, "spaceAfter": 8, "alignment": "left", "firstLineIndent": 0},
        },
    }
    validate_config(conf)
    return conf


def parse_inline_formatting(text: str) -> list[Dict[str, Any]]:
    """
    Parse inline formatting from text and return a list of segments.
    Each segment has: text, bold, italic, code, link (url or None)
    
    Supports: **bold**, *italic*, `code`, [text](url)
    """
    segments = []
    
    # Combined pattern for all inline formatting
    # Order matters: bold (**) before italic (*) to avoid conflicts
    # Using non-greedy matching and proper escaping
    pattern = r'(\*\*(.+?)\*\*)|(?<!\*)(\*([^*]+?)\*)(?!\*)|(`([^`]+)`)|(\[([^\]]+)\]\(([^)]+)\))'
    
    last_end = 0
    for match in re.finditer(pattern, text):
        # Add any text before this match as plain text
        if match.start() > last_end:
            plain_text = text[last_end:match.start()]
            if plain_text:
                segments.append({
                    'text': plain_text,
                    'bold': False,
                    'italic': False,
                    'code': False,
                    'link': None
                })
        
        # Determine which group matched
        if match.group(1):  # Bold: **text**
            segments.append({
                'text': match.group(2),
                'bold': True,
                'italic': False,
                'code': False,
                'link': None
            })
        elif match.group(3):  # Italic: *text*
            segments.append({
                'text': match.group(4),
                'bold': False,
                'italic': True,
                'code': False,
                'link': None
            })
        elif match.group(5):  # Code: `text`
            segments.append({
                'text': match.group(6),
                'bold': False,
                'italic': False,
                'code': True,
                'link': None
            })
        elif match.group(7):  # Link: [text](url)
            segments.append({
                'text': match.group(8),
                'bold': False,
                'italic': False,
                'code': False,
                'link': match.group(9)
            })
        
        last_end = match.end()
    
    # Add any remaining text after the last match
    if last_end < len(text):
        remaining = text[last_end:]
        if remaining:
            segments.append({
                'text': remaining,
                'bold': False,
                'italic': False,
                'code': False,
                'link': None
            })
    
    # If no formatting found, return the whole text as plain
    if not segments:
        segments.append({
            'text': text,
            'bold': False,
            'italic': False,
            'code': False,
            'link': None
        })
    
    return segments


def add_hyperlink(paragraph, url: str, text: str, style_config: Dict[str, Any], global_config: Dict[str, Any]):
    """Add a hyperlink to a paragraph."""
    # Create the hyperlink element
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    run = paragraph.add_run(text)
    apply_run_fmt(run, style_config, global_config)

    # Override with link styling
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.underline = True

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)


def add_formatted_runs(paragraph, text: str, base_style: Dict[str, Any], global_config: Dict[str, Any], code_style: Dict[str, Any] = None) -> None:
    """Add text with inline formatting to a paragraph."""
    segments = parse_inline_formatting(text)
    
    for segment in segments:
        if segment['link']:
            # Add hyperlink
            add_hyperlink(paragraph, segment['link'], segment['text'], base_style, global_config)
        else:
            run = paragraph.add_run(segment['text'])
            
            # Apply base style first
            if segment['code'] and code_style:
                apply_run_fmt(run, code_style, global_config)
            else:
                apply_run_fmt(run, base_style, global_config)
            
            # Override with inline formatting
            if segment['bold']:
                run.bold = True
            if segment['italic']:
                run.italic = True
            if segment['code']:
                # Use monospace font for inline code
                run.font.name = 'Courier New'
                _ensure_east_asia_font(run, global_config.get('baseFontCn', 'SimSun'), 'Courier New')


def add_heading(doc: Document, text: str, level: int, conf: Dict[str, Any]) -> None:
    style = conf["styles"].get(f"h{level}", conf["styles"]["h1"])
    p = doc.add_paragraph()
    apply_paragraph_fmt(p, style)
    code_style = conf["styles"].get("code", {})
    add_formatted_runs(p, text.strip(), style, conf["global"], code_style)


def add_body(doc: Document, text: str, conf: Dict[str, Any]) -> None:
    style = conf["styles"]["body"]
    p = doc.add_paragraph()
    apply_paragraph_fmt(p, style)
    
    # Check if text starts with markdown formatting tag
    # If so, remove the first line indent
    stripped = text.lstrip()
    # Matches: **Bold**, *Italic*, `Code`, [Link], ![Image]
    if re.match(r"^(\*\*|\*|`|\[|!\[)", stripped):
        p.paragraph_format.first_line_indent = 0
        
    code_style = conf["styles"].get("code", {})
    add_formatted_runs(p, text, style, conf["global"], code_style)


def add_quote(doc: Document, text: str, conf: Dict[str, Any]) -> None:
    style = conf["styles"]["quote"]
    p = doc.add_paragraph()
    apply_paragraph_fmt(p, style)
    # 引用缩进
    try:
        p.paragraph_format.left_indent = Inches(0.25)
    except Exception:
        pass
    code_style = conf["styles"].get("code", {})
    add_formatted_runs(p, text, style, conf["global"], code_style)


def add_list_item(doc: Document, text: str, ordered: bool, conf: Dict[str, Any]) -> None:
    style = conf["styles"]["body"]
    p = doc.add_paragraph()
    apply_paragraph_fmt(p, style)
    # 使用内置样式简化项目符号/编号
    try:
        p.style = "List Number" if ordered else "List Bullet"
    except Exception:
        pass
    code_style = conf["styles"].get("code", {})
    add_formatted_runs(p, text, style, conf["global"], code_style)


def add_code_block(doc: Document, lines: list[str], conf: Dict[str, Any]) -> None:
    style = conf["styles"]["code"]
    for line in lines:
        p = doc.add_paragraph()
        apply_paragraph_fmt(p, style)
        run = p.add_run(line if line.strip() else " ")
        apply_run_fmt(run, style, conf["global"])


def add_horizontal_rule(doc: Document) -> None:
    """Add a horizontal rule (line divider) to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # Create a bottom border on the paragraph to simulate a horizontal line
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_border(cell, border_color: str = "000000", border_size: int = 4) -> None:
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_cell_shading(cell, fill_color: str) -> None:
    """Set background shading for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_color)
    tcPr.append(shading)


def _split_table_row(line: str) -> list[str]:
    cells: list[str] = []
    current: list[str] = []
    i = 0
    while i < len(line):
        ch = line[i]
        if ch == "\\" and i + 1 < len(line) and line[i + 1] == "|":
            current.append("|")
            i += 2
            continue
        if ch == "|":
            cells.append("".join(current).strip())
            current = []
            i += 1
            continue
        current.append(ch)
        i += 1
    cells.append("".join(current).strip())
    return cells


def _parse_alignment_markers(cells: list[str]) -> Optional[list[Optional[str]]]:
    alignments: list[Optional[str]] = []
    for cell in cells:
        trimmed = cell.strip()
        if not trimmed or "-" not in trimmed or set(trimmed) - set("-:"):
            return None
        if trimmed.startswith(":") and trimmed.endswith(":"):
            alignments.append("center")
        elif trimmed.startswith(":"):
            alignments.append("left")
        elif trimmed.endswith(":"):
            alignments.append("right")
        else:
            alignments.append(None)
    return alignments


def parse_gfm_table(lines: list[str]) -> Optional[Dict[str, Any]]:
    """
    Parse GFM table lines into a 2D list of cell contents.
    Returns None if the lines don't form a valid GFM table.
    """
    if len(lines) < 2:
        return None
    
    separator_line = lines[1].strip()
    if separator_line.startswith("|"):
        separator_line = separator_line[1:]
    if separator_line.endswith("|"):
        separator_line = separator_line[:-1]
    separator_cells = _split_table_row(separator_line)
    alignments = _parse_alignment_markers(separator_cells)
    if alignments is None:
        return None
    
    rows = []
    for i, line in enumerate(lines):
        if i == 1:  # Skip separator line
            continue
        # Parse cells from the line
        line = line.strip()
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        cells = _split_table_row(line)
        if cells:
            rows.append(cells)
    
    if not rows:
        return None
    return {"rows": rows, "alignments": alignments}


def add_table(doc: Document, table_data: list[list[str]], conf: Dict[str, Any], alignments: Optional[list[Optional[str]]] = None) -> None:
    """Add a table to the document with proper borders and header styling."""
    if not table_data or not table_data[0]:
        return
    
    num_rows = len(table_data)
    num_cols = max(len(row) for row in table_data)
    
    # Create table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    style = conf["styles"].get("body", {})
    global_conf = conf.get("global", {})
    
    for row_idx, row_data in enumerate(table_data):
        row = table.rows[row_idx]
        for col_idx in range(num_cols):
            cell = row.cells[col_idx]
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            
            # Set cell borders
            _set_cell_border(cell)
            
            # Header row styling (first row)
            if row_idx == 0:
                _set_cell_shading(cell, "E5E7EB")  # Light gray background
            
            # Add text to cell
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(cell_text)
            apply_run_fmt(run, style, global_conf)

            if alignments and col_idx < len(alignments):
                align = _get_alignment(alignments[col_idx])
                if align is not None:
                    paragraph.alignment = align
            
            # Bold for header row
            if row_idx == 0:
                run.bold = True


def set_page_margins(doc: Document, margin_inch: float) -> None:
    for section in doc.sections:
        section.top_margin = Inches(margin_inch)
        section.bottom_margin = Inches(margin_inch)
        section.left_margin = Inches(margin_inch)
        section.right_margin = Inches(margin_inch)


def is_table_line(line: str) -> bool:
    """Check if a line looks like part of a GFM table."""
    stripped = line.strip()
    if not stripped:
        return False
    # Table lines contain pipes and are not code fences
    if '|' in stripped and not stripped.startswith('```'):
        return True
    return False


def is_table_separator(line: str) -> bool:
    """Check if a line is a GFM table separator (e.g., |---|---|)."""
    separator_pattern = r'^\s*\|?\s*[-:]+[-|\s:]*\|?\s*$'
    return bool(re.match(separator_pattern, line.strip()))


def convert(input_path: str, output_path: str, conf: Dict[str, Any]) -> None:
    """Convert Markdown file to Word document with proper error handling."""
    # Validate input file exists
    if not os.path.exists(input_path):
        raise FileError(
            "Input file not found",
            path=input_path
        )
    
    # Check output directory is writable
    output_dir = os.path.dirname(output_path) or '.'
    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir, exist_ok=True)
        except PermissionError as e:
            raise FileError(
                "Cannot create output directory",
                path=output_dir,
                details=str(e)
            )
    
    try:
        doc = Document()
    except Exception as e:
        raise DocxGenerationError(
            "Failed to create Word document",
            details=str(e)
        )
    
    try:
        margin_value = float(conf.get("global", {}).get("pageMargin", 1.0))
    except Exception as e:
        raise ConfigError("Invalid pageMargin value", details=str(e))
    set_page_margins(doc, margin_value)

    try:
        with open(input_path, "r", encoding="utf-8") as f:
            lines = f.read().splitlines()
    except PermissionError as e:
        raise FileError(
            "Permission denied reading input file",
            path=input_path,
            details=str(e)
        )
    except UnicodeDecodeError as e:
        raise ConversionError(
            "Failed to decode input file (expected UTF-8 encoding)",
            details=str(e)
        )

    in_code = False
    code_buf: list[str] = []
    in_table = False
    table_buf: list[str] = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        
        # Handle code blocks
        fence = re.match(r"^```", line)
        if fence:
            # Flush table buffer if we were in a table
            if in_table and table_buf:
                table_data = parse_gfm_table(table_buf)
                if table_data:
                    add_table(doc, table_data["rows"], conf, table_data["alignments"])
                table_buf = []
                in_table = False
            
            if in_code:
                add_code_block(doc, code_buf, conf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
            
        if in_code:
            code_buf.append(line)
            i += 1
            continue
        
        # Check for table start: current line has pipes and next line is separator
        if not in_table and is_table_line(line) and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            in_table = True
            table_buf = [line]
            i += 1
            continue
        
        # Continue collecting table lines
        if in_table:
            if is_table_line(line) or is_table_separator(line):
                table_buf.append(line)
                i += 1
                continue
            else:
                # End of table
                table_data = parse_gfm_table(table_buf)
                if table_data:
                    add_table(doc, table_data["rows"], conf, table_data["alignments"])
                table_buf = []
                in_table = False
                # Don't increment i, process current line normally
        
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            add_heading(doc, text, level, conf)
            i += 1
            continue
        if re.match(r"^\s*>\s+(.*)$", line):
            text = re.sub(r"^\s*>\s+", "", line)
            add_quote(doc, text, conf)
            i += 1
            continue
        if re.match(r"^\s*-\s+(.*)$", line):
            text = re.sub(r"^\s*-\s+", "", line)
            add_list_item(doc, text, ordered=False, conf=conf)
            i += 1
            continue
        if re.match(r"^\s*\d+\.\s+(.*)$", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            add_list_item(doc, text, ordered=True, conf=conf)
            i += 1
            continue
        # Horizontal rule: ---, ***, ___
        if re.match(r"^\s*([-*_])\s*\1\s*\1\s*$", line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue
        if line.strip() == "":
            # 忽略空行，用户可在 Word 中自行调整
            i += 1
            continue
        add_body(doc, line, conf)
        i += 1
    
    # Flush any remaining buffers
    if in_code and code_buf:
        add_code_block(doc, code_buf, conf)
    if in_table and table_buf:
        table_data = parse_gfm_table(table_buf)
        if table_data:
            add_table(doc, table_data["rows"], conf, table_data["alignments"])

    try:
        doc.save(output_path)
    except PermissionError as e:
        raise FileError(
            "Permission denied writing output file",
            path=output_path,
            details=str(e)
        )
    except Exception as e:
        raise DocxGenerationError(
            "Failed to save Word document",
            details=str(e)
        )


def main():
    """Main entry point with comprehensive error handling."""
    parser = argparse.ArgumentParser(description="Markdown to Word (.docx) converter with style config")
    parser.add_argument("--input", "-i", required=True, help="输入的 Markdown 文件路径")
    parser.add_argument("--output", "-o", required=True, help="输出的 .docx 文件路径")
    parser.add_argument("--config", "-c", help="JSON 字符串形式的样式配置")
    parser.add_argument("--config-file", "-f", help="JSON 配置文件路径")
    args = parser.parse_args()
    
    try:
        conf = load_config(args)
        convert(args.input, args.output, conf)
    except FileError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(e.exit_code if hasattr(e, 'exit_code') else EXIT_FILE_NOT_FOUND)
    except ConfigError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(e.exit_code if hasattr(e, 'exit_code') else EXIT_CONFIG_ERROR)
    except ConversionError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(e.exit_code if hasattr(e, 'exit_code') else EXIT_MARKDOWN_PARSE_ERROR)
    except DocxGenerationError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(e.exit_code if hasattr(e, 'exit_code') else EXIT_DOCX_GENERATION_ERROR)
    except Exception as e:
        print(f"Unexpected error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
